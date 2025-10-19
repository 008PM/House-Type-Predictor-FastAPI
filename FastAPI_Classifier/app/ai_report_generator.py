"""
AI-Powered Erläuterungsbericht Generator
Uses Claude Sonnet 4.5 for intelligent content generation
"""

from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, Any, Optional, List
import pandas as pd
import os
from dotenv import load_dotenv

load_dotenv()


class AIReportGenerator:
    """
    Generates professional Erläuterungsberichte using Claude AI
    """
    
    def __init__(self, project_name: str, location: str, 
                 project_type: str, federal_state: str):
        self.project_name = project_name
        self.location = location
        self.project_type = project_type
        self.federal_state = federal_state
        
        # Data
        self.room_book_data = None
        self.cost_data = None
        self.room_summary = None
        
        # AI Client
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY not found in environment")
        
        self.claude = Anthropic(api_key=api_key)
        
        print(f"✓ AI Report Generator initialized for: {project_name}")
    
    def load_room_book(self, file) -> bool:
        """Load and analyze room book"""
        try:
            df = pd.read_excel(file.file if hasattr(file, 'file') else file)
            self.room_book_data = df
            
            # Analyze room data for AI context
            self.room_summary = {
                "total_rooms": len(df),
                "total_area": df['area_m2'].sum() if 'area_m2' in df.columns else None,
                "room_types": df['room_type'].value_counts().to_dict() if 'room_type' in df.columns else {}
            }
            
            print(f"✓ Loaded room book: {len(df)} rooms")
            return True
        except Exception as e:
            print(f"✗ Error loading room book: {e}")
            return False
    
    def load_cost_estimate(self, file) -> bool:
        """Load cost estimate data"""
        try:
            df = pd.read_excel(file.file if hasattr(file, 'file') else file)
            self.cost_data = df
            print(f"✓ Loaded cost estimate: {len(df)} rows")
            return True
        except Exception as e:
            print(f"✗ Error loading costs: {e}")
            return False
    
    def generate_report(self) -> Dict[str, Any]:
        """
        Generate complete report with AI-generated content
        """
        print("\n🤖 Starting AI-powered report generation...")
        
        # Create project context for Claude
        context = self._build_project_context()
        
        # Generate all sections with AI
        sections = {
            "A1_allgemeines": self._generate_section_a1_ai(context),
            "A2_erschliessung": self._generate_section_a2_ai(context),
            "A3_kg410": self._generate_section_kg410_ai(context),
            "A4_kg420": self._generate_section_kg420_ai(context),
            "A5_kg434": self._generate_section_kg434_ai(context),
            "A6_kg430": self._generate_section_kg430_ai(context),
            "A7_kg440": self._generate_section_kg440_ai(context),
            "A8_kg470": self._generate_section_kg470_ai(context),
            "A9_kg480": self._generate_section_kg480_ai(context),
            "B_costs": self._generate_cost_summary()
        }
        
        report = {
            "metadata": {
                "title": "Erläuterungsbericht zum Vorentwurf",
                "subtitle": "Technische Gebäudeausrüstung",
                "project_name": self.project_name,
                "location": self.location,
                "date": datetime.now().strftime("%d.%m.%Y"),
                "author": "BKW AI Planning Assistant (powered by Claude)"
            },
            "sections": sections
        }
        
        print("✓ Report generation complete!\n")
        return report
    
    def _build_project_context(self) -> str:
        """Build comprehensive context for Claude"""
        
        context = f"""
PROJEKT-KONTEXT FÜR ERLÄUTERUNGSBERICHT:

Projektname: {self.project_name}
Standort: {self.location}
Gebäudetyp: {self.project_type}
Bundesland: {self.federal_state}
"""
        
        # Add room book data if available
        if self.room_summary:
            context += f"""
GEBÄUDEDATEN:
- Anzahl Räume: {self.room_summary['total_rooms']}
- Gesamtfläche: {self.room_summary['total_area']:.1f} m² (geschätzt)
- Raumtypen: {', '.join([f"{k}: {v}" for k, v in list(self.room_summary['room_types'].items())[:5]])}
"""
        
        # Add cost data if available
        if self.cost_data is not None:
            context += f"\nKOSTENDATA: {len(self.cost_data)} Positionen verfügbar"
        
        return context
    
    def _call_claude(self, prompt: str, max_tokens: int = 2000) -> str:
        """
        Call Claude API with error handling
        """
        try:
            message = self.claude.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=max_tokens,
                temperature=0.7,
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            )
            
            # Extract text from response
            response_text = message.content[0].text
            
            # Log token usage
            print(f"  → Claude used {message.usage.input_tokens} input + {message.usage.output_tokens} output tokens")
            
            return response_text
            
        except Exception as e:
            print(f"✗ Claude API error: {e}")
            return f"[Fehler bei AI-Generierung: {str(e)}]"
    
    def _generate_section_a1_ai(self, context: str) -> Dict:
        """
        A.1 Allgemeines - AI-generated
        """
        print("🤖 Generating A.1 Allgemeines...")
        
        # A.1.1 Aufgabenstellung
        prompt_task = f"""{context}

Schreibe einen professionellen Abschnitt "Aufgabenstellung" für einen Erläuterungsbericht (Leistungsphase 2 nach HOAI).

Der Abschnitt soll enthalten:
1. Kurze Beschreibung des Bauvorhabens
2. Umfang der TGA-Planung
3. Liste der betroffenen Gewerke (KG 410, 420, 430, 440, 470, 480)
4. Hinweis auf Planungsgrundlagen

Stil: Sachlich, professionell, wie ein deutscher Ingenieurbericht
Länge: 250-300 Wörter
Format: Fließtext in deutscher Sprache

WICHTIG: Schreibe NUR den Textinhalt, keine Überschriften, keine Markdown-Formatierung."""

        task_text = self._call_claude(prompt_task)
        
        # A.1.3 Gebäude
        prompt_building = f"""{context}

Beschreibe das Gebäude für den Abschnitt "Gebäude" im Erläuterungsbericht.

Berücksichtige:
- Gebäudetyp ({self.project_type})
- Nutzungskonzept
- Konstruktion (allgemein)
- Besonderheiten für TGA-Planung relevant

Stil: Sachlich, technisch präzise
Länge: 150-200 Wörter
Format: Fließtext

WICHTIG: Schreibe NUR den Textinhalt, keine Überschriften."""

        building_text = self._call_claude(prompt_building)
        
        # A.1.5 GEG
        prompt_geg = f"""{context}

Schreibe einen Abschnitt zum Gebäudeenergiegesetz (GEG) für dieses Projekt.

Erwähne:
- Anwendbarkeit des GEG 2024
- Energetische Anforderungen
- Geplante Maßnahmen zur Erfüllung
- Primärenergiefaktor

Stil: Sachlich, normkonform
Länge: 120-150 Wörter

WICHTIG: Schreibe NUR den Textinhalt, keine Überschriften."""

        geg_text = self._call_claude(prompt_geg, max_tokens=800)
        
        return {
            "title": "A.1 Allgemeines",
            "subsections": {
                "A.1.1 Aufgabenstellung": task_text,
                "A.1.2 Lage": f"Das Projekt befindet sich in {self.location}.",
                "A.1.3 Gebäude": building_text,
                "A.1.5 Gebäudeenergiegesetz GEG": geg_text,
                "A.1.6 Relevante Normen und Vorschriften": self._get_standards_formatted()
            }
        }
    
    def _generate_section_a2_ai(self, context: str) -> Dict:
        """A.2 Öffentliche Erschließung - AI-generated"""
        print("🤖 Generating A.2 Öffentliche Erschließung...")
        
        prompt = f"""{context}

Schreibe den Abschnitt "KG 220 - Öffentliche Erschließung" für einen Erläuterungsbericht.

Unterabschnitte:
1. KG 221 - Abwasserentsorgung
2. KG 222 - Wasserversorgung  
3. KG 224 - Wärmeversorgung (Anschluss)
4. KG 225 - Stromversorgung (Anschluss)

Für jeden Unterabschnitt: 2-3 Sätze über die geplante Anbindung an öffentliche Netze.

Stil: Technisch präzise, sachlich
Gesamtlänge: 300-400 Wörter

Format: Strukturiere mit **Überschriften** für jeden Unterabschnitt."""

        content = self._call_claude(prompt, max_tokens=1500)
        
        return {
            "title": "A.2 KG 220 - Öffentliche Erschließung",
            "content": content
        }
    
    def _generate_section_kg410_ai(self, context: str) -> Dict:
        """A.3 KG 410 - Abwasser, Wasser, Gas - AI-generated"""
        print("🤖 Generating A.3 KG 410 Sanitäranlagen...")
        
        prompt = f"""{context}

Erstelle den Abschnitt "KG 410 - Abwasser-, Wasser- und Gasanlagen" für einen deutschen Erläuterungsbericht.

Unterabschnitte (jeweils technisches Konzept beschreiben):

**A.3.1 KG 411 - Schmutzwasseranlagen**
- Entwässerungssystem (Trennsystem/Mischsystem)
- Ableitung
- Besonderheiten

**A.3.2 KG 411 - Regenentwässerung**
- Dachent wässerung
- Rückhaltung/Versickerung
- Ableitung

**A.3.3 KG 412 - Trinkwasserversorgung**
- Kaltwasserversorgung
- Warmwasserversorgung
- Zirkulation
- Hygieneanforderungen nach Trinkwasserverordnung

Berücksichtige:
- Gebäudetyp: {self.project_type}
- Aktuelle Normen: DIN 1986, DIN 1988, DIN EN 806
- Trinkwasserhygiene VDI 6023

Stil: Technisch fundiert, wie in einem LP2-Bericht
Länge: 500-600 Wörter total
Format: Mit klaren Unterüberschriften (**fett**)"""

        content = self._call_claude(prompt, max_tokens=2500)
        
        return {
            "title": "A.3 KG 410 - Abwasser-, Wasser- und Gasanlagen",
            "content": content
        }
    
    def _generate_section_kg420_ai(self, context: str) -> Dict:
        """A.4 KG 420 - Wärmeversorgung - AI-generated"""
        print("🤖 Generating A.4 KG 420 Wärmeversorgung...")
        
        prompt = f"""{context}

Erstelle den Abschnitt "KG 420 - Wärmeversorgungsanlagen" für einen professionellen deutschen Erläuterungsbericht.

Unterabschnitte:

**A.4.1 KG 421 - Wärmeerzeugungsanlagen**
- Wahl der Wärmeerzeugung (Wärmepumpe, Fernwärme, Gas-Brennwert, etc.)
- Begründung der Systemwahl
- Auslegung nach DIN EN 12831
- Dimensionierung
- Spitzenlast-Abdeckung

**A.4.2 KG 421 - Zentrale Warmwasserbereitung**
- System zur Warmwasserbereitung
- Speicherkonzept
- Legionellenschutz

**A.4.3 KG 422 - Wärmeverteilnetze**
- Verteilsystem (2-Leiter, 4-Leiter)
- Temperaturniveaus (VL/RL)
- Hydraulik
- Dämmung

**A.4.4 KG 423 - Raumheizflächen**
- Typ der Heizflächen (Fußbodenheizung, Heizkörper, Konvektoren)
- Zuordnung zu Raumbereichen
- Regelungskonzept

Berücksichtige:
- Gebäudetyp: {self.project_type}
- Standort: {self.federal_state} (Klimazone)
- GEG 2024 Anforderungen
- Moderne, energieeffiziente Lösungen
- Normen: DIN EN 12831, DIN EN 12828, VDI 2035

Stil: Ingenieurmäßig, fundiert, entscheidungsbegründend
Länge: 600-700 Wörter
Format: Klar strukturiert mit **Unterüberschriften**"""

        content = self._call_claude(prompt, max_tokens=3000)
        
        return {
            "title": "A.4 KG 420 - Wärmeversorgungsanlagen",
            "content": content
        }
    
    def _generate_section_kg434_ai(self, context: str) -> Dict:
        """A.5 KG 434 - Kälte - AI-generated"""
        print("🤖 Generating A.5 KG 434 Kältetechnik...")
        
        prompt = f"""{context}

Schreibe den Abschnitt "KG 434 - Kältetechnische Anlagen".

Beschreibe:
- Kältebedarf (wo und warum)
- Kälteerzeugung (Kompressionskälte, Adsorption, etc.)
- Kälteverteilung
- Rückkühlung
- Kälteabgabe (Kühldecken, Kühlbalken, etc.)

Berücksichtige Gebäudetyp: {self.project_type}

Wenn für diesen Gebäudetyp üblich: Detailliertes Konzept
Wenn nicht üblich: Kurz erwähnen "ggf. dezentrale Split-Geräte für Serverräume"

Normen: VDI 2078 (Kühllast), DIN EN 378 (Kälteanlagen)

Länge: 250-350 Wörter
Stil: Technisch präzise"""

        content = self._call_claude(prompt, max_tokens=1500)
        
        return {
            "title": "A.5 KG 434 - Kältetechnische Anlagen",
            "content": content
        }
    
    def _generate_section_kg430_ai(self, context: str) -> Dict:
        """A.6 KG 430 - Lüftung - AI-generated"""
        print("🤖 Generating A.6 KG 430 Lüftungstechnik...")
        
        prompt = f"""{context}

Erstelle den Abschnitt "KG 430 - Lüftungstechnische Anlagen" für einen LP2 Erläuterungsbericht.

**A.6.1 Grundlagen**
- Lüftungsbedarf (Hygieneluft, Komfort)
- Normengrundlage (DIN 1946-6, DIN EN 16798)
- Luftqualitätskategorie

**A.6.2 RLT-Konzept**
- Anzahl und Art der RLT-Anlagen
- Zentral vs. Dezentral
- Funktionen (Heizen, Kühlen, Be-/Entfeuchten)
- Wärmerückgewinnung (WRG-Grad)
- Luftmengen
- Regelungskonzept
- Energieeffizienz (SFP-Klasse)

**A.6.3 Luftverteilung**
- Luftführung (Kanalführung, Schächte)
- Luftauslässe
- Luftdurchlässe

Berücksichtige:
- Gebäudetyp: {self.project_type}
- Moderne RLT-Technik mit WRG
- Normen: DIN 1946, DIN EN 13779, DIN EN 16798, VDI 6022

Länge: 600-700 Wörter
Format: Mit klaren **Unterüberschriften**
Stil: Technisch fundiert, entscheidungsbegründend"""

        content = self._call_claude(prompt, max_tokens=3000)
        
        return {
            "title": "A.6 KG 430 - Lüftungstechnische Anlagen",
            "content": content
        }
    
    def _generate_section_kg440_ai(self, context: str) -> Dict:
        """A.7 KG 440 - Elektro - AI-generated"""
        print("🤖 Generating A.7 KG 440 Elektroanlagen...")
        
        prompt = f"""{context}

Erstelle den Abschnitt "KG 440 - Elektroanlagen".

Unterabschnitte:

**A.7.1 Stromversorgung Allgemein**
- Netzanschluss
- Leistungsbedarf
- Versorgungssicherheit

**A.7.2 Niederspannungshauptverteilung (NSHV)**
- Standort
- Dimensionierung
- Unterverteilungen

**A.7.3 Niederspannungsinstallation**
- Installationssystem
- Verlegearten
- FI/LS-Schutz

**A.7.4 Beleuchtung**
- Beleuchtungskonzept (LED)
- Lichtsteuerung (Tageslicht, Präsenz)
- Beleuchtungsstärken nach DIN EN 12464-1
- Notbeleuchtung

**A.7.5 Blitzschutz und Erdung**
- Blitzschutzsystem nach DIN EN 62305
- Potentialausgleich

**A.7.6 Photovoltaik** (falls für {self.project_type} relevant)

Normen: DIN VDE 0100, DIN EN 12464-1, DIN EN 62305

Länge: 500-600 Wörter
Format: Mit **Unterüberschriften**
Stil: Normkonform, technisch präzise"""

        content = self._call_claude(prompt, max_tokens=2500)
        
        return {
            "title": "A.7 KG 440 - Elektroanlagen",
            "content": content
        }
    
    def _generate_section_kg470_ai(self, context: str) -> Dict:
        """A.8 KG 470 - Nutzungsspezifische Anlagen - AI-generated"""
        print("🤖 Generating A.8 KG 470 Nutzungsspezifische Anlagen...")
        
        prompt = f"""{context}

Schreibe den Abschnitt "KG 470 - Nutzungsspezifische Anlagen".

Für Gebäudetyp {self.project_type}, berücksichtige typische nutzungsspezifische Anlagen:

Büro: Feuerlöschanlage, ggf. Küchentechnik
Labor: Laborgasversorgung, Sicherheitseinrichtungen, Abzüge
Krankenhaus: Medizinische Gase, Vakuum, Druckluft, Feuerlöschanlage
Schule: Feuerlöschanlage, ggf. Küchentechnik

**A.8.1 KG 474 - Feuerlöschanlagen** (wenn relevant)
- Sprinkleranlage
- Wandhydranten
- Konzept

**A.8.2 Weitere** (je nach Typ)

Wenn für den Gebäudetyp nicht relevant: Kurz schreiben "Für dieses Projekt nicht vorgesehen"

Länge: 200-300 Wörter
Stil: Sachlich, sicherheitsorientiert"""

        content = self._call_claude(prompt, max_tokens=1500)
        
        return {
            "title": "A.8 KG 470 - Nutzungsspezifische Anlagen",
            "content": content
        }
    
    def _generate_section_kg480_ai(self, context: str) -> Dict:
        """A.9 KG 480 - Gebäudeautomation - AI-generated"""
        print("🤖 Generating A.9 KG 480 Gebäudeautomation...")
        
        prompt = f"""{context}

Erstelle den Abschnitt "KG 480 - Gebäudeautomation (GA)".

**A.9.1 GA-Konzept**
- Automationsgrad nach DIN EN ISO 16484
- DDC-System (Direkte Digitale Regelung)
- Kommunikationsprotokoll (BACnet, KNX, etc.)

**A.9.2 Funktionen**
- Einzelraumregelung
- Anlagenregelung (RLT, Heizung, Kühlung)
- Energiemanagement
- Visualisierung
- Fernzugriff

**A.9.3 Integration**
- Schnittstellen zu TGA-Anlagen
- Alarmierung
- Zeitprogramme

Gebäudetyp: {self.project_type}

Normen: DIN EN ISO 16484, VDI 3814

Länge: 400-500 Wörter
Format: Mit **Unterüberschriften**
Stil: Technisch, zukunftsorientiert"""

        content = self._call_claude(prompt, max_tokens=2000)
        
        return {
            "title": "A.9 KG 480 - Gebäudeautomation",
            "content": content
        }
    
    def _generate_cost_summary(self) -> Dict:
        """B. Kostenschätzung - with or without data"""
        
        if self.cost_data is not None:
            content = """Die detaillierte Kostenschätzung ist der separaten Excel-Datei zu entnehmen.

Die Kosten wurden ermittelt auf Basis von:
- BKI Baukosten (Baupreisindex aktuell)
- Vergleichsprojekten
- Herstellerangaben für Hauptkomponenten
- Erfahrungswerten

Die Kostenschätzung erfolgt gemäß DIN 276 mit einer Genauigkeit von ±30%."""
        else:
            content = """Die Kostenschätzung erfolgt auf Grundlage von:
- BKI Baukosteninformation (aktuelle Indices)
- Erfahrungswerten aus vergleichbaren Projekten  
- Herstellerangaben für Großkomponenten
- Pauschalansätzen für Kleinteile

**Kostenschätzung nach DIN 276** (vorläufig, Genauigkeit ±30%):

- KG 410 Sanitär: [zu ermitteln] €
- KG 420 Wärmeversorgung: [zu ermitteln] €
- KG 434 Kälte: [zu ermitteln] €
- KG 430 Lüftung: [zu ermitteln] €
- KG 440 Elektro: [zu ermitteln] €
- KG 470 Nutzungsspez.: [zu ermitteln] €
- KG 480 Gebäudeautomation: [zu ermitteln] €

**Summe TGA (KG 400)**: [zu ermitteln] € (netto)

Die detaillierte Kostenschätzung wird in der nächsten Bearbeitungsphase erstellt."""

        return {
            "title": "B. Kostenschätzung",
            "content": content
        }
    
    def _get_standards_formatted(self) -> str:
        """Get relevant standards as formatted text"""
        
        standards = [
            "DIN EN 12831 - Heizlastberechnung",
            "DIN EN 16798 - Energetische Bewertung von Gebäuden - Lüftung",
            "DIN 1946 - Raumlufttechnik",
            "DIN 1988 - Technische Regeln für Trinkwasser-Installationen",
            "DIN 1986 - Entwässerungsanlagen für Gebäude und Grundstücke",
            "VDI 2078 - Kühllastberechnung",
            "VDI 6023 - Hygiene in Trinkwasser-Installationen",
            "DIN VDE 0100 - Errichten von Niederspannungsanlagen",
            "DIN EN 12464-1 - Beleuchtung von Arbeitsstätten",
            "DIN EN ISO 16484 - Gebäudeautomation",
        ]
        
        # Add state-specific building code
        state_codes = {
            "Bayern": "Bayerische Bauordnung (BayBO)",
            "Baden-Württemberg": "Landesbauordnung Baden-Württemberg (LBO)",
            "Nordrhein-Westfalen": "Bauordnung NRW (BauO NRW)",
            "Hessen": "Hessische Bauordnung (HBO)",
            "Berlin": "Bauordnung Berlin (BauO Bln)",
        }
        
        if self.federal_state in state_codes:
            standards.append(state_codes[self.federal_state])
        
        # Add project-type specific standards
        if self.project_type == "laboratory":
            standards.extend([
                "DIN 12924 - Laboreinrichtungen",
                "DIN 1946-7 - Raumlufttechnik in Laboratorien"
            ])
        elif self.project_type == "hospital":
            standards.extend([
                "DIN 1946-4 - Raumlufttechnik in Krankenhäusern",
                "DIN VDE 0100-710 - Medizinisch genutzte Bereiche"
            ])
        
        return "\n".join([f"• {std}" for std in standards])
    
    def export_docx(self, report: Dict[str, Any]) -> str:
        """
        Export report as professional DOCX with proper formatting
        """
        print("\n📄 Exporting to DOCX...")
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # TITLE PAGE
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(report["metadata"]["project_name"])
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(report["metadata"]["title"])
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.bold = True
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        subtitle2_para = doc.add_paragraph()
        subtitle2_run = subtitle2_para.add_run(report["metadata"]["subtitle"])
        subtitle2_run.font.size = Pt(14)
        subtitle2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Standort: ").bold = True
        meta_para.add_run(report["metadata"]["location"])
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        date_para = doc.add_paragraph()
        date_para.add_run(f"Datum: ").bold = True
        date_para.add_run(report["metadata"]["date"])
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        author_para = doc.add_paragraph()
        author_para.add_run(f"Erstellt mit: ").bold = True
        author_para.add_run(report["metadata"]["author"])
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # TABLE OF CONTENTS
        toc_heading = doc.add_heading("Inhaltsverzeichnis", 1)
        doc.add_paragraph("[Inhaltsverzeichnis wird in Microsoft Word automatisch erstellt]")
        doc.add_paragraph("In Word: Referenzen → Inhaltsverzeichnis → Automatisches Verzeichnis")
        
        doc.add_page_break()
        
        # SECTIONS
        for section_key, section_data in report["sections"].items():
            if isinstance(section_data, dict):
                # Main section heading
                doc.add_heading(section_data.get("title", section_key), 1)
                
                # Content or subsections
                if "content" in section_data:
                    # Simple content
                    self._add_formatted_content(doc, section_data["content"])
                
                elif "subsections" in section_data:
                    # Subsections
                    for sub_key, sub_content in section_data["subsections"].items():
                        doc.add_heading(sub_key, 2)
                        self._add_formatted_content(doc, sub_content)
                
                doc.add_page_break()
        
        # Save
        filename = f"Erlaeuterungsbericht_{self.project_name.replace(' ', '_').replace('/', '_')}.docx"
        output_path = f"/tmp/{filename}"
        doc.save(output_path)
        
        print(f"✓ DOCX saved: {filename}")
        return output_path
    
    def _add_formatted_content(self, doc, content: str):
        """
        Add content with basic Markdown-style formatting support
        """
        # Split by paragraphs
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Check for bold headings (**text**)
            if para_text.startswith('**') and para_text.endswith('**'):
                # It's a heading
                heading_text = para_text.strip('*')
                doc.add_heading(heading_text, 3)
            elif '**' in para_text:
                # Mixed formatting - need to parse
                para = doc.add_paragraph()
                parts = para_text.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        # Normal text
                        para.add_run(part)
                    else:
                        # Bold text
                        para.add_run(part).bold = True
            else:
                # Simple paragraph
                if para_text.startswith('- ') or para_text.startswith('• '):
                    # Bullet point
                    doc.add_paragraph(para_text.lstrip('- •'), style='List Bullet')
                else:
                    doc.add_paragraph(para_text)
    
    def export_markdown(self, report: Dict[str, Any]) -> str:
        """Export as Markdown"""
        print("\n📝 Exporting to Markdown...")
        
        md_content = f"""# {report["metadata"]["project_name"]}

## {report["metadata"]["title"]}
### {report["metadata"]["subtitle"]}

**Standort:** {report["metadata"]["location"]}  
**Datum:** {report["metadata"]["date"]}  
**Erstellt mit:** {report["metadata"]["author"]}

---

## Inhaltsverzeichnis

[Automatisches Inhaltsverzeichnis]

---

"""
        
        for section_key, section_data in report["sections"].items():
            if isinstance(section_data, dict):
                md_content += f"\n## {section_data.get('title', section_key)}\n\n"
                
                if "content" in section_data:
                    md_content += f"{section_data['content']}\n\n"
                
                elif "subsections" in section_data:
                    for sub_key, sub_content in section_data["subsections"].items():
                        md_content += f"\n### {sub_key}\n\n{sub_content}\n\n"
                
                md_content += "\n---\n"
        
        filename = f"Erlaeuterungsbericht_{self.project_name.replace(' ', '_')}.md"
        output_path = f"/tmp/{filename}"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        print(f"✓ Markdown saved: {filename}")
        return output_path


# Quick test
if __name__ == "__main__":
    print("Testing AI Report Generator...")
    
    generator = AIReportGenerator(
        project_name="Test Bürogebäude",
        location="München, Bayern",
        project_type="office",
        federal_state="Bayern"
    )
    
    report = generator.generate_report()
    docx_path = generator.export_docx(report)
    
    print(f"\n✅ Test complete! Check: {docx_path}")